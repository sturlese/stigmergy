"""Deterministic converters (the HANDS): pdf, sheet, docx, office, text.

Each one returns faithful text; the judgment (quality, representation, cleanup) belongs to the agent.
"""
import os
import re
import subprocess
import tempfile

import httpx
from docx import Document
from openpyxl import load_workbook

EXT_METHOD = {
    ".pdf": "pdf",
    ".xlsx": "sheet", ".xlsm": "sheet", ".xls": "sheet", ".csv": "sheet", ".tsv": "sheet",
    ".docx": "docx",
    # .ods goes through LibreOffice (Gotenberg), NOT openpyxl — openpyxl cannot read OpenDocument
    # spreadsheets.
    ".pptx": "office", ".ppt": "office", ".doc": "office", ".odt": "office", ".odp": "office",
    ".ods": "office", ".rtf": "office",
    ".md": "text", ".txt": "text", ".json": "text",
}

SHEET_MAX_ROWS = 5000     # hard cap on rows read per sheet
SHEET_SAMPLE_ROWS = 25    # compact profile the agent sees


def method_for_ext(ext: str) -> str:
    return EXT_METHOD.get(ext.lower(), "text")


# `-layout` can hard-wrap a long token across a line break, so a credential can arrive
# already split. `librarian.gates` scans adjacent line pairs rejoined to compensate —
# changing this flag changes what that gate can see.
def _pdftotext(path: str) -> str:
    r = subprocess.run(["pdftotext", "-layout", path, "-"], capture_output=True, text=True)
    if r.returncode != 0:
        raise RuntimeError(f"pdftotext rc={r.returncode}: {r.stderr[:500]}")
    return r.stdout


def _office_to_pdf(path: str) -> bytes:
    url = os.environ.get("GOTENBERG_URL", "http://gotenberg:3000")  # call-time read, never at import
    with open(path, "rb") as f:
        files = {"files": (os.path.basename(path), f.read())}
    resp = httpx.post(f"{url}/forms/libreoffice/convert", files=files, timeout=180)
    if resp.status_code != 200:
        raise RuntimeError(f"gotenberg {resp.status_code}: {resp.text[:300]}")
    return resp.content


def _docx_to_md(path: str) -> str:
    doc = Document(path)
    parts = []
    for p in doc.paragraphs:
        t = p.text.strip()
        if not t:
            continue
        style = (p.style.name or "").lower() if p.style else ""
        if "heading" in style:
            parts.append(f"## {t}")
        else:
            parts.append(t)
    for tbl in doc.tables:
        for row in tbl.rows:
            cells = [c.text.strip().replace("\n", " ") for c in row.cells]
            parts.append("| " + " | ".join(cells) + " |")
    return "\n\n".join(parts)


def _render_table(rs) -> str:
    """Every row padded to the WIDEST row, not to row 0: a sheet that opens with a one-cell title
    above the grid would otherwise render as a one-column table with full-width rows hanging off
    it. A rectangular grid is unaffected (max width IS row 0's) — and it has to be, because this
    text is what `librarian.gates` scans for secrets, line by line."""
    def cell(c):
        return c.replace("|", "\\|").replace("\n", " ").strip()

    width = max(len(r) for r in rs)

    def row(r):
        return "| " + " | ".join([cell(c) for c in r] + [""] * (width - len(r))) + " |"

    return "\n".join([row(rs[0]), "| " + " | ".join(["---"] * width) + " |"]
                     + [row(r) for r in rs[1:]])


def _xls_rows(path: str) -> list:
    """Read legacy .xls (old BIFF) with xlrd -> [(sheet_name, rows)]. openpyxl can't read .xls."""
    import xlrd
    book = xlrd.open_workbook(path)
    out = []
    for sh in book.sheets():
        rows = []
        # The cap counts rows KEPT, not rows scanned — the shape `_xlsx_rows` and `_csv_rows` use.
        # Capping the scan instead loses every real row of a sparse sheet whose grid starts below
        # the cap, which is judgment about the data in the module that promises none.
        for ri in range(sh.nrows):
            cells = []
            for v in sh.row_values(ri):
                if v is None or v == "":
                    cells.append("")
                elif isinstance(v, float) and v.is_integer():
                    cells.append(str(int(v)))  # 5000.0 -> "5000"
                else:
                    cells.append(str(v))
            if any(c.strip() for c in cells):
                rows.append(cells)
            if len(rows) >= SHEET_MAX_ROWS:
                break
        out.append((sh.name, rows))
    return out


def _xlsx_rows(path: str) -> list:
    wb = load_workbook(path, read_only=True, data_only=True)
    out = []
    for name in wb.sheetnames:
        ws = wb[name]
        rows = []
        for row in ws.iter_rows(values_only=True):
            if row and any(c is not None and str(c).strip() != "" for c in row):
                rows.append(["" if c is None else str(c) for c in row])
            if len(rows) >= SHEET_MAX_ROWS:
                break
        out.append((name, rows))
    return out


def _csv_rows(path: str, delim: str) -> list:
    import csv
    rows = []
    with open(path, newline="", encoding="utf-8", errors="replace") as f:
        for r in csv.reader(f, delimiter=delim):
            cells = ["" if c is None else str(c) for c in r]
            if any(c.strip() for c in cells):
                rows.append(cells)
            if len(rows) >= SHEET_MAX_ROWS:
                break
    return [("Sheet1", rows)]


def sheet_rows(path: str) -> list:
    """Full parsed grid per sheet: [(sheet_name, rows)] with rows capped at SHEET_MAX_ROWS."""
    ext = os.path.splitext(path)[1].lower()
    if ext == ".xls":
        return _xls_rows(path)
    if ext == ".csv":
        return _csv_rows(path, ",")
    if ext == ".tsv":
        return _csv_rows(path, "\t")
    return _xlsx_rows(path)


def _sheet_profile(path: str) -> str:
    """Compact per-sheet profile: dimensions + a sample of rows. A reading aid, never the
    record — the full grid stays in the original file the evidence plane keeps."""
    sheets = sheet_rows(path)
    parts = []
    for name, rows in sheets:
        if not rows:
            continue
        sample = rows[: SHEET_SAMPLE_ROWS + 1]
        parts.append(
            f"### {name} ({len(rows)} rows, {max(len(r) for r in rows)} cols)\n\n{_render_table(sample)}"
            + (f"\n\n_(+{len(rows) - len(sample)} more rows in the original file)_" if len(rows) > len(sample) else "")
        )
    return "\n\n".join(parts)


VISION_OCR_PROMPT = (
    "Transcribe this document FAITHFULLY and COMPLETELY to text/Markdown. Keep tables as tables. "
    "Mark illegible passages as [illegible]. Do not summarize, invent or interpret — ONLY transcribe "
    "the content exactly as it appears."
)

DEFAULT_VISION_MODEL = "gemini-3-flash-preview"

# The prefixed form's bounds. Pages past the ceiling are not transcribed — the text SAYS where
# it was cut, because a transcription that silently stops reads as complete. `-scale-to` bounds
# the OUTPUT box where a fixed DPI does not: a legal max-MediaBox page (200×200 in) at 150 dpi
# is a raster bomb (30000×30000 px) hiding inside the drive door's 25 MB cap. Both subprocesses
# and the model request carry their own clocks, so a hostile document costs a bounded slice of
# the item's lease, never the whole of it.
MAX_VISION_PAGES = 40
MAX_VISION_RASTER_PX = 2000     # the long side of one rasterized page
PDF_RASTER_TIMEOUT_S = 120
PDFINFO_TIMEOUT_S = 30
VISION_CALL_TIMEOUT_S = 240


def vision_config_error() -> str | None:
    """Why `vision_extract` cannot run, or `None` when it can — the question
    `librarian.processing` asks BEFORE paying for a call, worded for the refusal it lands in.

    A provider-prefixed model checks its provider's key against the kernel's one table: a KNOWN
    prefix with no key is the misconfiguration "requeue" can never fix, so it must read as
    unconfigured with the variable named. An UNKNOWN prefix is configured by naming itself — the
    same warn-don't-refuse posture as the librarian's preflight, so the gate stays
    provider-agnostic. The bare Gemini form is configured by its own key."""
    from stigmergy.kernel.settings import PROVIDER_KEY_ENV, provider_of

    model = os.environ.get("VISION_MODEL") or DEFAULT_VISION_MODEL
    provider = provider_of(model)
    if not provider:
        if os.environ.get("GEMINI_API_KEY"):
            return None
        return ("no vision OCR is configured — the operator can set GEMINI_API_KEY (for the "
                "default Gemini model) or a provider-prefixed VISION_MODEL with its provider's "
                "key, and requeue")
    key_env = PROVIDER_KEY_ENV.get(provider)
    if key_env and not os.environ.get(key_env):
        return (f"VISION_MODEL names a {provider}: model but {key_env} is not set — the "
                f"operator can set it and requeue")
    return None


def vision_configured() -> bool:
    """`vision_config_error() is None`, for the caller that only branches."""
    return vision_config_error() is None


def _pdf_page_pngs(path: str, cap: int = MAX_VISION_PAGES) -> list[bytes]:
    """The first `cap` pages as PNG bytes, via poppler's pdftoppm — the same package
    `_pdftotext` ships in, so this path owes the image no new toolchain. pdftoppm zero-pads its
    page numbers, so the lexical sort IS page order (pinned against the real binary by
    `tests/kernel/test_converters.py`)."""
    with tempfile.TemporaryDirectory() as td:
        prefix = os.path.join(td, "page")
        try:
            r = subprocess.run(["pdftoppm", "-png", "-scale-to", str(MAX_VISION_RASTER_PX),
                                "-l", str(cap), path, prefix],
                               capture_output=True, text=True, timeout=PDF_RASTER_TIMEOUT_S)
        except subprocess.TimeoutExpired as ex:
            raise RuntimeError(
                f"pdftoppm exceeded its {PDF_RASTER_TIMEOUT_S}s budget rasterizing this "
                f"PDF") from ex
        if r.returncode != 0:
            raise RuntimeError(f"pdftoppm rc={r.returncode}: {r.stderr[:500]}")
        pages = []
        for name in sorted(os.listdir(td)):
            with open(os.path.join(td, name), "rb") as f:
                pages.append(f.read())
    if not pages:
        raise RuntimeError("pdftoppm produced no pages — the PDF may be corrupt or empty")
    return pages


def _pdf_page_count(path: str) -> int:
    """`pdfinfo`'s Pages figure, `0` when it cannot say — read only to decide the cut note, and
    AFTER the model call, so "cannot say" (missing binary, timeout, bad rc) must degrade to `0`
    rather than discard a paid transcription."""
    try:
        r = subprocess.run(["pdfinfo", path], capture_output=True, text=True,
                           timeout=PDFINFO_TIMEOUT_S)
    except (OSError, subprocess.TimeoutExpired):
        return 0
    if r.returncode != 0:
        return 0
    m = re.search(r"^Pages:\s+(\d+)", r.stdout, re.MULTILINE)
    return int(m.group(1)) if m else 0


def _vision_extract_pydantic(model: str, path: str, agent_builder=None) -> dict:
    """The prefixed form: page IMAGES to any pydantic-ai model — the pages are decided by CODE,
    so the OCR means the same thing whatever the provider serves it. `agent_builder` is the
    injectable seam a test drives a scripted model through; production passes none.

    The agent is built BEFORE the rasterization: a typo'd model must fail before the expensive
    step, not after it. The returned dict carries `pages`/`truncated` as DATA beside the in-text
    cut note — the note is for the filing model (which needs to judge completeness), the fields
    are for code and logs, which must never parse a sentence out of untrusted-adjacent text."""
    from pydantic_ai import Agent, BinaryContent

    from stigmergy.kernel.llm import build_model
    from stigmergy.kernel.usage_repair import ensure_usage_extraction_repaired

    # Installed HERE beside this module's own `Agent(...)` — every agent-construction site
    # installs the repair (idempotent), and `tests/kernel/test_usage_repair.py` holds that rule
    # textually, on purpose.
    ensure_usage_extraction_repaired()
    # Through `build_model` for one reason: the `model_override` seam (#81) applies here too, so
    # a test can drive this path with a scripted model. A prefixed string comes back verbatim.
    resolved, _ = build_model(model)
    agent = agent_builder(model) if agent_builder else Agent(
        resolved, model_settings={"timeout": VISION_CALL_TIMEOUT_S})
    pages = _pdf_page_pngs(path)
    parts = [VISION_OCR_PROMPT] + [BinaryContent(data=png, media_type="image/png")
                                   for png in pages]
    result = agent.run_sync(parts)
    text = str(result.output or "")
    truncated = False
    if len(pages) == MAX_VISION_PAGES:
        total = _pdf_page_count(path)
        truncated = total == 0 or total > MAX_VISION_PAGES
        if truncated:
            text += (f"\n\n[the worker cut this document here: it is longer than the "
                     f"{MAX_VISION_PAGES}-page OCR ceiling, so what you have is its opening "
                     f"and not the whole of it]")
    return {"method": "vision", "text": text, "model": model,
            "pages": len(pages), "truncated": truncated}


def vision_extract(path: str, agent_builder=None) -> dict:
    """OCR a scanned/visual PDF. TWO forms of `$VISION_MODEL`, the same convention CLEAN_MODEL
    and ANSWER_MODEL follow: a bare name (the default) is Gemini's native-PDF single call and
    requires `GEMINI_API_KEY`; a provider-prefixed pydantic-ai id
    ("openrouter:qwen/qwen3-vl-8b-instruct") sends poppler-rasterized page images instead, and
    authenticates with that provider's own key. Lazy SDK imports on both forms."""
    model = os.environ.get("VISION_MODEL") or DEFAULT_VISION_MODEL
    if ":" in model:
        return _vision_extract_pydantic(model, path, agent_builder)
    key = os.environ.get("GEMINI_API_KEY")
    if not key:
        raise RuntimeError("method=vision requires GEMINI_API_KEY (Google AI Studio); not set in the environment")
    from google import genai
    from google.genai import types
    client = genai.Client(api_key=key)
    if os.path.getsize(path) <= 14 * 1024 * 1024:
        # INLINE bytes: no Files API -> no filename header (non-ASCII filenames break the Files
        # API's ASCII header encoding). Gemini accepts native PDF inline up to ~20MB.
        with open(path, "rb") as f:
            part = types.Part.from_bytes(data=f.read(), mime_type="application/pdf")
    else:
        # large PDF -> Files API, but through a copy with an ASCII name (avoids the header issue)
        with open(path, "rb") as src, tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as tmp:
            tmp.write(src.read())
            tmpname = tmp.name
        try:
            part = client.files.upload(file=tmpname)
        finally:
            os.unlink(tmpname)
    r = client.models.generate_content(model=model, contents=[part, VISION_OCR_PROMPT])
    return {"method": "vision", "text": r.text or "", "model": model}   # model -> faithful provenance


def extract(path: str, method: str) -> dict:
    """Returns {method, text}. `text` is what the agent sees."""
    if method == "pdf":
        return {"method": "pdf", "text": _pdftotext(path)}
    if method == "office":
        pdf = _office_to_pdf(path)
        with tempfile.NamedTemporaryFile(suffix=".pdf", delete=True) as tmp:
            tmp.write(pdf)
            tmp.flush()
            return {"method": "office", "text": _pdftotext(tmp.name)}
    if method == "docx":
        return {"method": "docx", "text": _docx_to_md(path)}
    if method == "sheet":
        return {"method": "sheet", "text": _sheet_profile(path)}
    with open(path, encoding="utf-8", errors="replace") as f:
        return {"method": "text", "text": f.read()}
