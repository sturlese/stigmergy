"""Deterministic converters: routing, table rendering, sheet profile, fail-fast paths.
No network and no real binaries: subprocess/httpx/genai are monkeypatched."""
import subprocess
import sys
import types

import pytest

from stigmergy.kernel import converters
from stigmergy.kernel.converters import _csv_rows, _render_table, _sheet_profile, extract, method_for_ext


def test_method_for_ext_routing():
    assert method_for_ext(".PDF") == "pdf"
    assert method_for_ext(".xlsx") == "sheet"
    assert method_for_ext(".docx") == "docx"
    assert method_for_ext(".pptx") == "office"
    assert method_for_ext(".ods") == "office"   # LibreOffice route: openpyxl cannot read .ods
    assert method_for_ext(".md") == "text"
    assert method_for_ext(".weird") == "text"


def test_render_table_escapes_pipes_and_newlines():
    md = _render_table([["a|b", "h2"], ["x\ny", "z"]])
    assert "a\\|b" in md
    assert "x y" in md
    assert md.splitlines()[1] == "| --- | --- |"


def _md_width(line: str) -> int:
    """Cells in one rendered markdown row. `| a | b |` splits into 4 pipe-separated parts, two of
    them the empty ends."""
    return len(line.split("|")) - 2


def test_a_ragged_grid_renders_every_row_at_the_widest_row(tmp_path):
    """OLD BEHAVIOUR: header and separator were 1 cell wide while the data rows were 3 — a table
    whose whole width came from `rs[0]`.

    Real sheets open with a one-cell title row above the grid, so a 20-column export rendered as a
    1-column table with 20-column rows hanging off it: not valid markdown, and unreadable to the
    agent that is handed this text as the whole of a dropped spreadsheet.
    """
    p = tmp_path / "ragged.csv"
    p.write_text("Quarterly report\nmetric,q1,q2\narr,1000,1200\n")

    text = extract(str(p), "sheet")["text"]

    table = [ln for ln in text.splitlines() if ln.startswith("|")]
    assert [_md_width(ln) for ln in table] == [3, 3, 3, 3]
    assert table[1] == "| --- | --- | --- |"
    assert "| arr | 1000 | 1200 |" in table


def test_a_rectangular_grid_renders_exactly_as_it_always_did(tmp_path):
    """The benign twin, pinned as a golden string. This text feeds `librarian.gates`' secrets
    scan, which matches within a line — so a change to what these converters emit changes what
    that gate can see. A rectangular grid (every real sheet that has no title row) must come out
    byte-identical to before the ragged-row fix."""
    p = tmp_path / "rect.csv"
    p.write_text("metric,value\narr,1000\n")

    text = extract(str(p), "sheet")["text"]

    assert text == ("### Sheet1 (2 rows, 2 cols)\n\n"
                    "| metric | value |\n"
                    "| --- | --- |\n"
                    "| arr | 1000 |")


def test_csv_rows_skips_blank_lines(tmp_path):
    p = tmp_path / "d.csv"
    p.write_text("h1,h2\n\n1,2\n,,\n")
    (name, rows), = _csv_rows(str(p), ",")
    assert name == "Sheet1"
    assert rows == [["h1", "h2"], ["1", "2"]]


def test_sheet_profile_samples_and_counts(tmp_path):
    p = tmp_path / "big.csv"
    lines = ["col1,col2"] + [f"v{i},w{i}" for i in range(100)]
    p.write_text("\n".join(lines))
    profile = _sheet_profile(str(p))
    assert "(101 rows, 2 cols)" in profile
    assert "+75 more rows" in profile
    assert "v99" not in profile   # beyond the sample


def test_sheet_profile_xlsx(tmp_path):
    from openpyxl import Workbook
    wb = Workbook()
    ws = wb.active
    ws.title = "KPIs"
    ws.append(["metric", "value"])
    ws.append(["arr", 1000])
    p = tmp_path / "kpis.xlsx"
    wb.save(p)
    profile = _sheet_profile(str(p))
    assert "### KPIs (2 rows, 2 cols)" in profile
    assert "| arr | 1000 |" in profile


# ── .xls: the cap counts rows KEPT, the way its xlsx/csv twins do ──────────────────────────────
class FakeXlrdSheet:
    """xlrd's `Sheet` seam, exactly as `_xls_rows` reaches for it: `name`, `nrows`, `ncols`,
    `cell_value(r, c)` and the `row_values(r)` xlrd derives from them. This fake stands in for the
    EXTERNAL library (BIFF parsing), never for our own logic — the row selection under test is
    entirely `_xls_rows`'s."""

    def __init__(self, name, grid):
        self.name = name
        self._grid = grid
        self.nrows = len(grid)
        self.ncols = max((len(r) for r in grid), default=0)

    def cell_value(self, r, c):
        row = self._grid[r]
        return row[c] if c < len(row) else ""

    def row_values(self, r):
        return [self.cell_value(r, c) for c in range(self.ncols)]


def _xls_book(monkeypatch, sheets):
    import xlrd
    monkeypatch.setattr(xlrd, "open_workbook",
                        lambda path: types.SimpleNamespace(sheets=lambda: sheets))


def test_xls_rows_keeps_real_rows_that_sit_past_the_first_cap_worth_of_blanks(monkeypatch):
    """OLD BEHAVIOUR: `[]` — every real row lost, silently.

    The cap was applied to rows SCANNED (`range(min(sh.nrows, SHEET_MAX_ROWS))`) while the xlsx
    and csv readers apply it to rows KEPT. A sparse legacy .xls — a padded export, a grid pushed
    down by leading blank rows — therefore lost real data that an identical .xlsx keeps, in the
    module whose whole contract is faithful text with no judgment.
    """
    grid = [["", ""] for _ in range(converters.SHEET_MAX_ROWS)]
    grid += [["metric", "value"], ["arr", "1000"], ["nrr", "112"]]
    _xls_book(monkeypatch, [FakeXlrdSheet("Legacy", grid)])

    (name, rows), = converters._xls_rows("/legacy.xls")

    assert name == "Legacy"
    assert rows == [["metric", "value"], ["arr", "1000"], ["nrr", "112"]]


def test_xls_rows_still_stops_at_the_cap_on_a_dense_sheet(monkeypatch):
    """The benign twin: the cap is still a cap. A sheet with more real rows than SHEET_MAX_ROWS
    keeps exactly the first SHEET_MAX_ROWS of them — the fix moved what the cap counts, it did not
    remove the bound that keeps one dropped spreadsheet from becoming the worker's whole memory."""
    grid = [[f"r{i}", "x"] for i in range(converters.SHEET_MAX_ROWS + 10)]
    _xls_book(monkeypatch, [FakeXlrdSheet("Dense", grid)])

    (_name, rows), = converters._xls_rows("/dense.xls")

    assert len(rows) == converters.SHEET_MAX_ROWS
    assert rows[0] == ["r0", "x"]
    assert rows[-1] == [f"r{converters.SHEET_MAX_ROWS - 1}", "x"]


def test_xls_rows_renders_whole_floats_as_integers(monkeypatch):
    """xlrd hands every number back as a float; `5000.0` in a cell is `5000` to a reader. Pinned
    beside the cap tests because both live in the same loop."""
    _xls_book(monkeypatch, [FakeXlrdSheet("Nums", [[5000.0, 1.5, None]])])

    (_name, rows), = converters._xls_rows("/nums.xls")

    assert rows == [["5000", "1.5", ""]]


def test_extract_text(tmp_path):
    p = tmp_path / "note.md"
    p.write_text("hello")
    assert extract(str(p), "text") == {"method": "text", "text": "hello"}


def test_extract_pdf_uses_pdftotext(monkeypatch):
    def fake_run(cmd, capture_output, text):
        assert cmd[0] == "pdftotext"
        return types.SimpleNamespace(returncode=0, stdout="pdf text", stderr="")
    monkeypatch.setattr(subprocess, "run", fake_run)
    assert extract("/x.pdf", "pdf") == {"method": "pdf", "text": "pdf text"}


def test_pdftotext_failure_raises(monkeypatch):
    monkeypatch.setattr(subprocess, "run",
                        lambda *a, **k: types.SimpleNamespace(returncode=1, stdout="", stderr="boom"))
    with pytest.raises(RuntimeError, match="pdftotext"):
        extract("/x.pdf", "pdf")


def test_extract_office_roundtrips_through_gotenberg(monkeypatch, tmp_path):
    src = tmp_path / "deck.pptx"
    src.write_bytes(b"fake")
    monkeypatch.setattr(converters.httpx, "post",
                        lambda url, files, timeout: types.SimpleNamespace(status_code=200, content=b"%PDF"))
    monkeypatch.setattr(converters, "_pdftotext", lambda path: "converted text")
    assert extract(str(src), "office") == {"method": "office", "text": "converted text"}


def test_office_gotenberg_error_raises(monkeypatch, tmp_path):
    src = tmp_path / "deck.pptx"
    src.write_bytes(b"fake")
    monkeypatch.setattr(converters.httpx, "post",
                        lambda url, files, timeout: types.SimpleNamespace(status_code=500, text="err"))
    with pytest.raises(RuntimeError, match="gotenberg"):
        extract(str(src), "office")


def test_docx_extraction(tmp_path):
    from docx import Document
    doc = Document()
    doc.add_heading("Section", level=1)
    doc.add_paragraph("Body paragraph.")
    p = tmp_path / "m.docx"
    doc.save(p)
    res = extract(str(p), "docx")
    assert "## Section" in res["text"]
    assert "Body paragraph." in res["text"]


def test_vision_fails_fast_without_key(monkeypatch):
    monkeypatch.delenv("GEMINI_API_KEY", raising=False)
    monkeypatch.delenv("VISION_MODEL", raising=False)   # the bare default form is under test
    with pytest.raises(RuntimeError, match="GEMINI_API_KEY"):
        converters.vision_extract("/nonexistent.pdf")


def test_vision_extract_inline_small_pdf(monkeypatch, tmp_path):
    """Wires a fake google.genai module: small PDFs go inline, response text is returned."""
    p = tmp_path / "scan.pdf"
    p.write_bytes(b"%PDF fake")
    monkeypatch.setenv("GEMINI_API_KEY", "fake-key")

    calls = {}

    class FakeModels:
        def generate_content(self, model, contents):
            calls["model"] = model
            calls["contents"] = contents
            return types.SimpleNamespace(text="ocr result")

    class FakeClient:
        def __init__(self, api_key):
            calls["api_key"] = api_key
            self.models = FakeModels()
            self.files = types.SimpleNamespace(upload=lambda file: "uploaded")

    fake_types = types.SimpleNamespace(
        Part=types.SimpleNamespace(from_bytes=lambda data, mime_type: ("part", mime_type)))
    fake_genai = types.ModuleType("google.genai")
    fake_genai.Client = FakeClient
    fake_genai.types = fake_types
    fake_google = types.ModuleType("google")
    fake_google.genai = fake_genai
    monkeypatch.setitem(sys.modules, "google", fake_google)
    monkeypatch.setitem(sys.modules, "google.genai", fake_genai)
    monkeypatch.setitem(sys.modules, "google.genai.types", fake_types)

    res = converters.vision_extract(str(p))
    assert res["method"] == "vision"
    assert res["text"] == "ocr result"
    assert res["model"]                      # provenance present
    assert calls["contents"][0][1] == "application/pdf"   # inline part, not Files API


# ── the two-form vision dispatch: bare = Gemini (above), prefixed = pydantic-ai over pages ──────
class FakeVisionAgent:
    """The `agent_builder` seam's double: records what the prefixed path sends, answers once."""

    def __init__(self, output="transcribed text"):
        self.output = output
        self.parts = None

    def run_sync(self, parts):
        self.parts = parts
        return types.SimpleNamespace(output=self.output)


def test_vision_prefixed_model_sends_rasterized_pages_to_the_pydantic_path(monkeypatch):
    """OLD BEHAVIOUR: `vision_extract` was welded to the Gemini SDK, so OCR could not run on any
    other provider. A provider-prefixed VISION_MODEL now takes the pydantic-ai road: page images
    rasterized by CODE, one transcription back, GEMINI_API_KEY never consulted."""
    monkeypatch.setenv("VISION_MODEL", "openrouter:qwen/qwen3-vl-8b-instruct")
    monkeypatch.delenv("GEMINI_API_KEY", raising=False)
    monkeypatch.setattr(converters, "_pdf_page_pngs",
                        lambda path, cap=converters.MAX_VISION_PAGES: [b"png-1", b"png-2"])
    agent = FakeVisionAgent()

    res = converters.vision_extract("/scan.pdf", agent_builder=lambda model: agent)

    assert res == {"method": "vision", "text": "transcribed text",
                   "model": "openrouter:qwen/qwen3-vl-8b-instruct",
                   "pages": 2, "truncated": False, "usage": None}
    assert agent.parts[0] == converters.VISION_OCR_PROMPT
    assert [p.data for p in agent.parts[1:]] == [b"png-1", b"png-2"]
    assert {p.media_type for p in agent.parts[1:]} == {"image/png"}


def test_vision_prefixed_reports_its_token_usage_as_data(monkeypatch):
    """The pass costs money and the caller prices it (issue #110): the framework's RunUsage is
    normalized to plain in/out counts. Absent usage stays None — zero would read as free."""
    monkeypatch.setenv("VISION_MODEL", "openrouter:qwen/qwen3-vl-8b-instruct")
    monkeypatch.setattr(converters, "_pdf_page_pngs",
                        lambda path, cap=converters.MAX_VISION_PAGES: [b"png-1"])

    class _WithUsage(FakeVisionAgent):
        def run_sync(self, parts):
            self.parts = parts
            return types.SimpleNamespace(output=self.output,
                                         usage=types.SimpleNamespace(input_tokens=1200,
                                                                     output_tokens=340))

    res = converters.vision_extract("/scan.pdf", agent_builder=lambda model: _WithUsage())
    assert res["usage"] == {"input_tokens": 1200, "output_tokens": 340}


def test_vision_prefixed_says_where_it_cut_a_document_over_the_page_ceiling(monkeypatch):
    """A transcription that silently stops reads as complete — the cap must speak."""
    monkeypatch.setenv("VISION_MODEL", "openrouter:qwen/qwen3-vl-8b-instruct")
    cap = converters.MAX_VISION_PAGES
    monkeypatch.setattr(converters, "_pdf_page_pngs", lambda path, cap=cap: [b"p"] * cap)
    monkeypatch.setattr(converters, "_pdf_page_count", lambda path: cap + 5)

    res = converters.vision_extract("/big.pdf", agent_builder=lambda model: FakeVisionAgent())

    assert "cut this document here" in res["text"]
    assert str(cap) in res["text"]
    assert res["truncated"] is True and res["pages"] == cap


def test_vision_prefixed_exactly_at_the_ceiling_is_not_called_cut(monkeypatch):
    """The benign twin: a document of exactly MAX_VISION_PAGES pages was transcribed WHOLE, and
    calling it cut would teach readers to distrust complete transcriptions."""
    monkeypatch.setenv("VISION_MODEL", "openrouter:qwen/qwen3-vl-8b-instruct")
    cap = converters.MAX_VISION_PAGES
    monkeypatch.setattr(converters, "_pdf_page_pngs", lambda path, cap=cap: [b"p"] * cap)
    monkeypatch.setattr(converters, "_pdf_page_count", lambda path: cap)

    res = converters.vision_extract("/exact.pdf", agent_builder=lambda model: FakeVisionAgent())

    assert "cut this document here" not in res["text"]
    assert res["truncated"] is False


def test_pdf_page_pngs_drives_pdftoppm_with_the_documented_bounds(monkeypatch):
    """The rasterizer's contract: poppler's pdftoppm (the same package pdftotext ships in), PNG,
    the output box BOUNDED with -scale-to (a fixed DPI is a raster bomb on a max-MediaBox page),
    capped with -l, its own clock, pages back in page order."""
    seen = {}

    def fake_run(cmd, capture_output, text, timeout):
        seen["cmd"], seen["timeout"] = cmd, timeout
        prefix = cmd[-1]
        for i in (1, 2):
            with open(f"{prefix}-{i}.png", "wb") as f:
                f.write(f"png-{i}".encode())
        return types.SimpleNamespace(returncode=0, stdout="", stderr="")

    monkeypatch.setattr(subprocess, "run", fake_run)
    pages = converters._pdf_page_pngs("/x.pdf", cap=7)

    assert pages == [b"png-1", b"png-2"]
    assert seen["cmd"][:5] == ["pdftoppm", "-png", "-scale-to",
                               str(converters.MAX_VISION_RASTER_PX), "-l"]
    assert seen["cmd"][5:] == ["7", "/x.pdf", seen["cmd"][-1]]
    assert seen["timeout"] == converters.PDF_RASTER_TIMEOUT_S


def test_a_hung_rasterizer_degrades_instead_of_hanging_the_item(monkeypatch):
    """`subprocess.TimeoutExpired` becomes the same RuntimeError shape every converter failure
    takes, so `_with_vision_fallback` degrades to the text layer instead of crashing."""
    def hung(cmd, capture_output, text, timeout):
        raise subprocess.TimeoutExpired(cmd, timeout)

    monkeypatch.setattr(subprocess, "run", hung)
    with pytest.raises(RuntimeError, match="pdftoppm exceeded"):
        converters._pdf_page_pngs("/x.pdf")


def _pdf_with_pages(n: int) -> bytes:
    """A minimal, real, n-page PDF — hand-written 1.4 syntax, one text object per page, the same
    approach `tests/librarian/test_drive_processing_pg.py`'s `_tiny_pdf` takes for one page."""
    objects = []
    page_ids = list(range(3, 3 + 2 * n, 2))
    kids = " ".join(f"{i} 0 R" for i in page_ids)
    objects.append(b"<< /Type /Catalog /Pages 2 0 R >>")
    objects.append(f"<< /Type /Pages /Kids [{kids}] /Count {n} >>".encode())
    for k, pid in enumerate(page_ids, start=1):
        content = f"BT /F1 24 Tf 72 700 Td (Page {k}) Tj ET".encode()
        objects.append(
            f"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] /Contents {pid + 1} 0 R "
            f"/Resources << /Font << /F1 << /Type /Font /Subtype /Type1 "
            f"/BaseFont /Helvetica >> >> >> >>".encode())
        objects.append(b"<< /Length %d >>\nstream\n%s\nendstream" % (len(content), content))
    out = bytearray(b"%PDF-1.4\n")
    offsets = []
    for i, body in enumerate(objects, start=1):
        offsets.append(len(out))
        out += b"%d 0 obj\n%s\nendobj\n" % (i, body)
    xref_at = len(out)
    out += b"xref\n0 %d\n0000000000 65535 f \n" % (len(objects) + 1)
    for off in offsets:
        out += b"%010d 00000 n \n" % off
    out += (b"trailer\n<< /Size %d /Root 1 0 R >>\nstartxref\n%d\n%%%%EOF\n"
            % (len(objects) + 1, xref_at))
    return bytes(out)


def _require_pdftoppm():
    import shutil
    if shutil.which("pdftoppm"):
        return
    import os
    if os.environ.get("STIGMERGY_TEST_DSN"):
        pytest.fail("$STIGMERGY_TEST_DSN is set (CI mode) but pdftoppm is not on PATH — refusing "
                    "to skip the rasterizer-order test silently. poppler-utils ships it beside "
                    "pdftotext (see .github/workflows/ci.yml).")
    pytest.skip("pdftoppm not on PATH (brew install poppler) — the prefixed vision form's hand")


def test_the_real_pdftoppm_zero_pads_so_lexical_sort_is_page_order(tmp_path):
    """The load-bearing claim behind `_pdf_page_pngs`'s `sorted(...)`, pinned against the REAL
    binary — never fake what you are claiming to prove. Twelve pages is the threshold: at nine,
    an unpadded naming scheme still sorts correctly and this test would prove nothing."""
    _require_pdftoppm()
    path = tmp_path / "twelve.pdf"
    path.write_bytes(_pdf_with_pages(12))

    import subprocess as sp
    outdir = tmp_path / "out"
    outdir.mkdir()
    r = sp.run(["pdftoppm", "-png", "-scale-to", "200", str(path), str(outdir / "page")],
               capture_output=True, text=True, timeout=60)
    assert r.returncode == 0, r.stderr[:300]
    names = sorted(p.name for p in outdir.iterdir())
    assert names[0] == "page-01.png" and names[-1] == "page-12.png"
    assert names == [f"page-{i:02d}.png" for i in range(1, 13)]

    pages = converters._pdf_page_pngs(str(path))
    assert len(pages) == 12


def test_vision_config_error_knows_both_forms_and_names_the_right_variable(monkeypatch):
    """`librarian.processing` pays for a call only when this returns None. A KNOWN provider
    prefix with no key is unconfigured and the reason names THAT provider's variable (audit S3);
    an unknown prefix is configured by naming itself — the librarian preflight's own
    warn-don't-refuse posture; the bare form is configured by its key."""
    monkeypatch.delenv("GEMINI_API_KEY", raising=False)
    monkeypatch.delenv("VISION_MODEL", raising=False)
    monkeypatch.delenv("OPENROUTER_API_KEY", raising=False)
    assert "GEMINI_API_KEY" in converters.vision_config_error()

    monkeypatch.setenv("GEMINI_API_KEY", "fake")
    assert converters.vision_config_error() is None

    monkeypatch.delenv("GEMINI_API_KEY")
    monkeypatch.setenv("VISION_MODEL", "openrouter:qwen/qwen3-vl-8b-instruct")
    reason = converters.vision_config_error()
    assert "OPENROUTER_API_KEY" in reason and "GEMINI_API_KEY" not in reason
    assert converters.vision_configured() is False

    monkeypatch.setenv("OPENROUTER_API_KEY", "sk-or-fake")
    assert converters.vision_config_error() is None

    monkeypatch.setenv("VISION_MODEL", "custom-lab:experimental-vl")   # unknown prefix
    monkeypatch.delenv("OPENROUTER_API_KEY")
    assert converters.vision_config_error() is None


def test_sheet_profile_reports_the_widest_rows_column_count(tmp_path):
    """The profile header must count the WIDEST row, like the rendered table under it.
    OLD BEHAVIOUR: it reported `len(rows[0])`, so a one-cell title row made a 3-column
    grid read as "1 cols" above a correctly 3-column table."""
    p = tmp_path / "ragged.csv"
    p.write_text("title\na,b,c\nd,e,f\n", encoding="utf-8")
    profile = _sheet_profile(str(p))
    assert "3 cols" in profile
    assert "1 cols" not in profile
